

import docx
from io import BytesIO

class Stream(BytesIO):
    def __init__(self, data):
        self.data = data
        self.pos = 0

    def seek(self, offset, whence=0):
        if whence == 0:
            self.pos = offset
        elif whence == 1:
            self.pos += offset
        elif whence == 2:
            self.pos = len(self.data) - offset

    def tell(self):
        return self.pos

    def read(self, size=None):
        if size is None:
            buffer = self.data[self.pos:]
            self.pos = len(self.data)
        else:
            buffer = self.data[self.pos:self.pos + size]
            self.pos += size

        return buffer


def loadDoc(fname):
    """
    载入word数据
    :param fname: 文件路径
    :return: 文本数据
    """
    strs = ""
    file = docx.Document(Stream(fname))
    #file = docx.Document(fname)
    for p in file.paragraphs:
        strs += p.text.encode("utf-8").decode("utf-8") + "\n"
        # strs += p.text + "\n"
    return strs


file_path = './sample.docx'
with open(file_path, 'rb') as fi:
    content = fi.read()

ss = loadDoc(content)
print(ss)
